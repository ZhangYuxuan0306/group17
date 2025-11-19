from __future__ import annotations

import json
import uuid
from pathlib import Path
from typing import Iterable, List

try:
    from pypdf import PdfReader  # type: ignore
except ImportError:  # pragma: no cover - 兼容旧版本库
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except ImportError:
        PdfReader = None  # type: ignore

from .logger import get_logger
from .types import RawDocument

logger = get_logger(__name__)


def _read_text_file(path: Path) -> str:
    """读取纯文本或 Markdown 文件内容。"""
    return path.read_text(encoding="utf-8")


def _read_docx_file(path: Path) -> str:
    """提取 Word 文本，优先使用 docx2txt，缺失时回退到 python-docx。"""
    try:
        import docx2txt  # type: ignore

        return docx2txt.process(str(path))
    except ImportError:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ImportError(
                "读取 .docx 文件需要安装 docx2txt 或 python-docx。"
            ) from exc

        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _read_pdf_file(path: Path) -> str:
    """使用 pypdf 提取 PDF 页面的全部文本。"""
    if PdfReader is None:
        raise ImportError(
            "读取 .pdf 文件需要安装 pypdf 或 PyPDF2。请执行 `pip install pypdf`。"
        )

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


READERS = {
    ".txt": _read_text_file,
    ".md": _read_text_file,
    ".docx": _read_docx_file,
    ".pdf": _read_pdf_file,
}


def load_documents(document_dir: Path) -> List[RawDocument]:
    """读取目录中的所有合法文档，转换为 RawDocument 列表。"""

    documents: List[RawDocument] = []
    for path in sorted(document_dir.glob("**/*")):
        if not path.is_file():
            continue

        # 根据后缀匹配对应的解析器
        reader = READERS.get(path.suffix.lower())
        if not reader:
            logger.warning("Skipping unsupported file type: %s", path)
            continue

        try:
            text = reader(path)
        except Exception as exc:
            logger.error("Failed to read %s: %s", path, exc)
            continue

        # 基于文件路径生成可复现的 doc_id
        doc_id = uuid.uuid5(uuid.NAMESPACE_URL, str(path.resolve())).hex
        metadata = {
            "source_name": path.name,
            "source_path": str(path),
            "file_extension": path.suffix.lower(),
        }

        documents.append(
            RawDocument(
                doc_id=doc_id,
                text=text,
                source_path=str(path),
                metadata=metadata,
            )
        )

    # ✅ 新增：智能提示逻辑
    num_docs = len(documents)
    if num_docs == 0:
        print("No valid documents found in directory: %s", document_dir)
    elif num_docs == 1:
        print(
            "Loaded 1 document from %s: %s",
            document_dir,
            documents[0].metadata.get("source_name", "unknown"),
        )
    else:
        print("Loaded %s documents from %s", num_docs, document_dir)

    return documents

def export_documents(documents: Iterable[RawDocument], output_path: Path) -> None:
    """将原始文档元数据写入 JSONL，便于追踪来源。"""

    with output_path.open("w", encoding="utf-8") as f:
        for doc in documents:
            record = {
                "doc_id": doc.doc_id,
                "source_path": doc.source_path,
                "metadata": doc.metadata,
            }
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
